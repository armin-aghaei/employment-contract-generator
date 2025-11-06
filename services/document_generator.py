"""
Document Generator Service - Creates PDF and DOCX files from filled templates.

Converts the filled JSON template into professional legal documents.
"""
import os
import io
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from weasyprint import HTML  # Commented out for testing - requires system libraries
from azure.storage.blob import BlobServiceClient


class DocumentGenerator:
    """Generates PDF and DOCX documents from filled templates."""

    def __init__(self):
        """Initialize Azure Blob Storage client."""
        connection_string = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
        self.blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        self.container_name = os.getenv("AZURE_STORAGE_CONTAINER_NAME", "documents")

        # Ensure container exists
        try:
            self.blob_service_client.create_container(self.container_name)
        except Exception:
            # Container already exists
            pass

    def generate_docx(self, filled_template: Dict, session_id: str) -> tuple[str, int]:
        """
        Generate DOCX document from filled template.

        Args:
            filled_template: The filled template JSON structure
            session_id: Session ID for naming the file

        Returns:
            Tuple of (blob_url, file_size_bytes)
        """
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Recursively process template structure
        self._add_json_to_docx(doc, filled_template)

        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        file_size = len(docx_bytes.getvalue())

        # Upload to blob storage
        blob_name = f"{session_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        blob_client = self.blob_service_client.get_blob_client(
            container=self.container_name,
            blob=blob_name
        )

        blob_client.upload_blob(docx_bytes.getvalue(), overwrite=True)
        blob_url = blob_client.url

        return blob_url, file_size

    def generate_pdf(self, filled_template: Dict, session_id: str) -> tuple[str, int]:
        """
        Generate PDF document from filled template.

        Args:
            filled_template: The filled template JSON structure
            session_id: Session ID for naming the file

        Returns:
            Tuple of (blob_url, file_size_bytes)
        """
        # PDF generation temporarily disabled - requires system libraries
        raise NotImplementedError("PDF generation is temporarily disabled. Please use DOCX format instead.")

        # Convert JSON to HTML
        # html_content = self._json_to_html(filled_template)

        # Generate PDF from HTML
        # pdf_bytes = HTML(string=html_content).write_pdf()
        file_size = len(pdf_bytes)

        # Upload to blob storage
        blob_name = f"{session_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
        blob_client = self.blob_service_client.get_blob_client(
            container=self.container_name,
            blob=blob_name
        )

        blob_client.upload_blob(pdf_bytes, overwrite=True)
        blob_url = blob_client.url

        return blob_url, file_size

    def _add_json_to_docx(self, doc: Document, data: Any, level: int = 0) -> None:
        """
        Recursively add JSON content to DOCX document.

        Args:
            doc: Document object
            data: JSON data (dict, list, or primitive)
            level: Current nesting level for indentation
        """
        if isinstance(data, dict):
            for key, value in data.items():
                if key == "title" and level == 0:
                    # Main title
                    heading = doc.add_heading(value, level=0)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif key == "sections" and isinstance(value, list):
                    # Process sections
                    for section in value:
                        self._add_json_to_docx(doc, section, level + 1)
                elif key == "section_title":
                    # Section heading
                    doc.add_heading(value, level=min(level, 3))
                elif key == "content" and isinstance(value, list):
                    # Section content paragraphs
                    for paragraph_text in value:
                        p = doc.add_paragraph(paragraph_text)
                        p.paragraph_format.space_after = Pt(6)
                elif key == "clauses" and isinstance(value, list):
                    # Numbered clauses
                    for i, clause in enumerate(value, 1):
                        p = doc.add_paragraph(f"{i}. {clause}", style='List Number')
                        p.paragraph_format.space_after = Pt(6)
                elif key == "signature_block":
                    # Add spacing before signatures
                    doc.add_paragraph()
                    doc.add_paragraph()
                    self._add_json_to_docx(doc, value, level)
                elif isinstance(value, (dict, list)):
                    self._add_json_to_docx(doc, value, level + 1)
                else:
                    # Simple key-value pair
                    if value and str(value).strip():
                        p = doc.add_paragraph(f"{key}: {value}")
                        p.paragraph_format.space_after = Pt(3)

        elif isinstance(data, list):
            for item in data:
                self._add_json_to_docx(doc, item, level)

        else:
            # Primitive value
            if data and str(data).strip():
                p = doc.add_paragraph(str(data))
                p.paragraph_format.space_after = Pt(6)

    def _json_to_html(self, data: Dict) -> str:
        """
        Convert JSON template to HTML for PDF generation.

        Args:
            data: Filled template JSON

        Returns:
            HTML string
        """
        html_parts = ["""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {
                    size: letter;
                    margin: 1in;
                }
                body {
                    font-family: 'Times New Roman', serif;
                    font-size: 12pt;
                    line-height: 1.5;
                    color: #000;
                }
                h1 {
                    text-align: center;
                    font-size: 18pt;
                    font-weight: bold;
                    margin-bottom: 24pt;
                    text-transform: uppercase;
                }
                h2 {
                    font-size: 14pt;
                    font-weight: bold;
                    margin-top: 18pt;
                    margin-bottom: 12pt;
                }
                h3 {
                    font-size: 12pt;
                    font-weight: bold;
                    margin-top: 12pt;
                    margin-bottom: 6pt;
                }
                p {
                    margin-bottom: 6pt;
                    text-align: justify;
                }
                .clause {
                    margin-left: 0.5in;
                    margin-bottom: 6pt;
                }
                .signature-block {
                    margin-top: 48pt;
                    page-break-inside: avoid;
                }
                .signature-line {
                    margin-top: 36pt;
                    border-top: 1px solid #000;
                    width: 50%;
                    padding-top: 6pt;
                }
            </style>
        </head>
        <body>
        """]

        # Recursively process JSON
        html_parts.append(self._json_to_html_recursive(data))

        html_parts.append("""
        </body>
        </html>
        """)

        return "\n".join(html_parts)

    def _json_to_html_recursive(self, data: Any, level: int = 0) -> str:
        """
        Recursively convert JSON to HTML.

        Args:
            data: JSON data
            level: Current heading level

        Returns:
            HTML string
        """
        html_parts = []

        if isinstance(data, dict):
            for key, value in data.items():
                if key == "title" and level == 0:
                    html_parts.append(f"<h1>{self._escape_html(value)}</h1>")
                elif key == "sections" and isinstance(value, list):
                    for section in value:
                        html_parts.append(self._json_to_html_recursive(section, level + 1))
                elif key == "section_title":
                    heading_level = min(level + 1, 3)
                    html_parts.append(f"<h{heading_level}>{self._escape_html(value)}</h{heading_level}>")
                elif key == "content" and isinstance(value, list):
                    for paragraph in value:
                        html_parts.append(f"<p>{self._escape_html(paragraph)}</p>")
                elif key == "clauses" and isinstance(value, list):
                    for i, clause in enumerate(value, 1):
                        html_parts.append(f'<p class="clause">{i}. {self._escape_html(clause)}</p>')
                elif key == "signature_block":
                    html_parts.append('<div class="signature-block">')
                    html_parts.append(self._json_to_html_recursive(value, level))
                    html_parts.append('</div>')
                elif isinstance(value, (dict, list)):
                    html_parts.append(self._json_to_html_recursive(value, level + 1))
                else:
                    if value and str(value).strip():
                        html_parts.append(f"<p><strong>{self._escape_html(key)}:</strong> {self._escape_html(value)}</p>")

        elif isinstance(data, list):
            for item in data:
                html_parts.append(self._json_to_html_recursive(item, level))

        else:
            if data and str(data).strip():
                html_parts.append(f"<p>{self._escape_html(data)}</p>")

        return "\n".join(html_parts)

    def _escape_html(self, text: Any) -> str:
        """Escape HTML special characters."""
        return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
